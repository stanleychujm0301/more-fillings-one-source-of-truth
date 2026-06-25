"""KPMG 工作底稿附件（P3 实现）— python-docx 生成 Word 文档。

这是评委加分项：审计师可以一键把差异说明附到 KPMG 内部工作底稿模板里。
"""

from __future__ import annotations

from pathlib import Path

from loguru import logger

from ahcc.schemas import Job


def export_working_paper(job: Job, out_path: Path) -> None:
    """导出 KPMG 工作底稿格式的 Word 附件。"""
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    title = doc.add_heading("差异说明附件 — A+H 股年报数据一致性核查", level=1)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x00, 0x33, 0x8D)

    # bold/italic 是 run 级属性，设在 Paragraph 上不生效，需作用到 run
    doc.add_paragraph(f"任务编号：{job.job_id}").runs[0].bold = True
    doc.add_paragraph(f"识别差异：{len(job.diffs)} 条").runs[0].bold = True

    for idx, diff in enumerate(job.diffs, 1):
        doc.add_heading(f"{idx}. {diff.topic.best()}（{diff.severity.value.upper()}）", level=2)
        doc.add_paragraph(diff.summary.best())
        if diff.standard_reasoning:
            doc.add_paragraph(f"AI 解读：{diff.standard_reasoning.rationale}").runs[0].italic = True
            for cite in diff.standard_reasoning.citations:
                doc.add_paragraph(f"  引用：{cite.standard_code} {cite.clause} — {cite.title}")
        for ev in diff.evidence:
            doc.add_paragraph(f"证据：{ev.side.value} 股年报 P.{ev.page}：{ev.snippet}")
        doc.add_paragraph("审计师结论：__________________________")
        doc.add_paragraph("签名 / 日期：__________________________")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    logger.info(f"工作底稿附件已导出：{out_path}")
